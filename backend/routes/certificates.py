"""
Certificates routes - Certificate management and generation
Endpoints: 12
"""
from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Query
from fastapi.responses import FileResponse
from typing import List, Optional
from datetime import datetime
from dateutil.relativedelta import relativedelta
from pathlib import Path
import uuid
import shutil
import subprocess
import tempfile
import copy
import logging

from docx import Document
from docx.shared import Pt

from core import db, get_current_user, get_malaysia_time, TEMPLATE_DIR, CERTIFICATE_PDF_DIR
from models import User

from pydantic import BaseModel, Field, ConfigDict

logger = logging.getLogger(__name__)

# Paths
STATIC_DIR = Path(__file__).parent.parent / "static"
CERTIFICATE_DIR = STATIC_DIR / "certificates"
CERTIFICATE_DIR.mkdir(parents=True, exist_ok=True)

# ==================== FONT SETTINGS SYSTEM ====================

DEFAULT_FONT_SETTINGS = {
    "participant_name": {"font_size": 16, "max_lines": 1, "auto_fit": True, "bold": True},
    "ic_number": {"font_size": 16, "max_lines": 1, "auto_fit": False, "bold": False},
    "company_name": {"font_size": 12, "max_lines": 1, "auto_fit": True, "bold": False},
    "certificate_title": {"font_size": 16, "max_lines": 2, "auto_fit": True, "bold": True},
    "certificate_subtitle": {"font_size": 12, "max_lines": 1, "auto_fit": True, "bold": False},
    "dates": {"font_size": 10, "max_lines": 1, "auto_fit": True, "bold": False},
    "venue": {"font_size": 8, "max_lines": 2, "auto_fit": True, "bold": False},
    "certificate_number": {"font_size": 10, "max_lines": 1, "auto_fit": False, "bold": False},
    "top_margin": 80,
    "paragraph_spacing": 65,
}

PLACEHOLDER_SETTINGS_MAP = {
    "{{PARTICIPANT_NAME}}": "participant_name",
    "{{IC_NUMBER}}": "ic_number",
    "{{COMPANY_NAME}}": "company_name",
    "{{CERTIFICATE_TITLE}}": "certificate_title",
    "{{CERTIFICATE_SUBTITLE}}": "certificate_subtitle",
    "{{VENUE}}": "venue",
    "{{CERTIFICATE_NUMBER}}": "certificate_number",
}

AREA_WIDTHS = {
    "participant_name": 150,
    "ic_number": 150,
    "company_name": 150,
    "certificate_title": 350,
    "certificate_subtitle": 350,
    "dates": 150,
    "venue": 150,
    "certificate_number": 350,
}


def _estimate_text_width(text: str, font_pt: float, bold: bool = False) -> float:
    avg_char_width = font_pt * 0.6
    if bold:
        avg_char_width *= 1.15
    return len(text) * avg_char_width


def _auto_fit_font(text: str, area_width: float, font_size: float, max_lines: int, bold: bool, min_pt: float = 6.0) -> float:
    size = font_size
    while size > min_pt:
        est_width = _estimate_text_width(text, size, bold)
        lines_needed = est_width / area_width if area_width > 0 else 1
        if lines_needed <= max_lines:
            return size
        size -= 0.5
    return min_pt


async def _get_font_settings() -> dict:
    saved = await db.settings.find_one({"id": "certificate_font_settings"}, {"_id": 0})
    if saved:
        merged = {**DEFAULT_FONT_SETTINGS}
        for k, v in saved.items():
            if k not in ("id", "_id"):
                merged[k] = v
        return merged
    return dict(DEFAULT_FONT_SETTINGS)


def _replace_placeholders_in_doc(doc: Document, replacements: dict, font_settings: dict):
    same_date = replacements.pop("_same_date", False)

    for paragraph in doc.paragraphs:
        para_text = paragraph.text

        if "{{CERTIFICATE_SUBTITLE}}" in para_text and not replacements.get("{{CERTIFICATE_SUBTITLE}}"):
            for run in paragraph.runs:
                run.text = ""
            continue

        if "{{VALIDITY_START}}" in para_text and not replacements.get("{{VALIDITY_END}}"):
            for run in paragraph.runs:
                run.text = ""
            continue

        is_date_line = "{{TRAINING_DATE}}" in para_text or "{{END_DATE}}" in para_text

        for run in paragraph.runs:
            if same_date and is_date_line and run.text.strip() == "-":
                run.text = ""
                continue

            for placeholder, value in replacements.items():
                if placeholder not in run.text:
                    continue
                run.text = run.text.replace(placeholder, value)

                settings_key = PLACEHOLDER_SETTINGS_MAP.get(placeholder)
                if settings_key and settings_key in font_settings:
                    fs = font_settings[settings_key]
                    if isinstance(fs, dict):
                        target_size = fs.get("font_size", 12)
                        if fs.get("auto_fit"):
                            target_size = _auto_fit_font(
                                value, AREA_WIDTHS.get(settings_key, 150),
                                target_size, fs.get("max_lines", 1), fs.get("bold", False),
                            )
                        run.font.size = Pt(target_size)

        if is_date_line:
            date_fs = font_settings.get("dates", {})
            date_size = date_fs.get("font_size", 10) if isinstance(date_fs, dict) else 10
            if isinstance(date_fs, dict) and date_fs.get("auto_fit"):
                full_text = "".join(r.text for r in paragraph.runs).strip()
                if full_text:
                    date_size = _auto_fit_font(
                        full_text, AREA_WIDTHS.get("dates", 150),
                        date_size, date_fs.get("max_lines", 1), False,
                    )
            for run in paragraph.runs:
                if run.text.strip():
                    run.font.size = Pt(date_size)


def _format_date_display(date_str: str) -> str:
    """Convert date string to display format (e.g., '25 January 2026')."""
    try:
        for fmt in ("%Y-%m-%d", "%Y-%m-%dT%H:%M:%S", "%d/%m/%Y", "%d-%m-%Y"):
            try:
                dt = datetime.strptime(date_str.split("T")[0] if "T" in date_str else date_str, fmt)
                return dt.strftime("%d %B %Y")
            except ValueError:
                continue
        return date_str
    except Exception:
        return date_str


async def _generate_cert_number() -> str:
    """Generate certificate number in format MDDRC/COA/YYYY/MM/00001."""
    now = get_malaysia_time()
    year = now.strftime("%Y")
    month = now.strftime("%m")
    prefix = f"MDDRC/COA/{year}/{month}/"
    count = await db.certificates.count_documents({"certificate_number": {"$regex": f"^MDDRC/COA/{year}/{month}/"}})
    return f"{prefix}{str(count + 1).zfill(5)}"


async def _build_replacements(session: dict, participant: dict, certificate_number: str) -> dict:
    """Build the placeholder→value mapping from session & participant data."""
    # Get program info
    program = None
    if session.get("program_id"):
        program = await db.programs.find_one({"id": session["program_id"]}, {"_id": 0})

    # Get company info
    company_name = session.get("company_name", "")
    if not company_name and session.get("company_id"):
        company = await db.companies.find_one({"id": session["company_id"]}, {"_id": 0})
        company_name = company.get("name", "") if company else ""

    # Certificate title/subtitle from program
    cert_title = ""
    cert_subtitle = ""
    if program:
        cert_title = program.get("certificate_title") or program.get("name", "")
        cert_subtitle = program.get("certificate_subtitle", "") or ""

    # Dates
    start_date = _format_date_display(session.get("start_date", ""))
    end_date = _format_date_display(session.get("end_date", ""))

    # If same date, use single date for display
    # The template has: "On {{TRAINING_DATE}}-{{END_DATE}}"
    # If dates are same, we set END_DATE to empty and TRAINING_DATE to the full date
    # and we'll handle the dash in the replacement
    same_date = (start_date == end_date)

    # Validity
    validity_start = end_date
    validity_end = ""
    if session.get("cert_show_validity", False):
        try:
            months = session.get("cert_validity_months", 24)
            end_dt_str = session.get("end_date", "")
            for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
                try:
                    end_dt = datetime.strptime(end_dt_str.split("T")[0], fmt)
                    validity_end_dt = end_dt + relativedelta(months=months)
                    validity_end = validity_end_dt.strftime("%d %B %Y")
                    break
                except ValueError:
                    continue
        except Exception:
            validity_end = ""

    return {
        "{{PARTICIPANT_NAME}}": participant.get("full_name", ""),
        "{{IC_NUMBER}}": participant.get("id_number", ""),
        "{{COMPANY_NAME}}": company_name,
        "{{CERTIFICATE_TITLE}}": cert_title,
        "{{CERTIFICATE_SUBTITLE}}": cert_subtitle,
        "{{TRAINING_DATE}}": start_date,
        "{{END_DATE}}": "" if same_date else end_date,
        "{{VENUE}}": session.get("location", ""),
        "{{VALIDITY_START}}": validity_start,
        "{{VALIDITY_END}}": validity_end,
        "{{CERTIFICATE_NUMBER}}": certificate_number,
        "_same_date": same_date,  # internal flag
    }


def _docx_to_pdf(docx_path: Path, output_dir: Path) -> Path:
    """Convert .docx to single-page .pdf using LibreOffice headless.
    Trims to page 1 if overflow occurs (common with complex templates)."""
    lo_path = "/usr/bin/libreoffice"
    result = subprocess.run(
        [lo_path, "--headless", "--convert-to", "pdf", "--outdir", str(output_dir), str(docx_path)],
        capture_output=True, text=True, timeout=60
    )
    if result.returncode != 0:
        logger.error(f"LibreOffice conversion failed: {result.stderr}")
        raise RuntimeError(f"PDF conversion failed: {result.stderr}")
    pdf_path = output_dir / (docx_path.stem + ".pdf")
    if not pdf_path.exists():
        raise RuntimeError("PDF file was not created by LibreOffice")

    # Trim to single page if overflow occurred
    from pypdf import PdfReader, PdfWriter
    reader = PdfReader(str(pdf_path))
    if len(reader.pages) > 1:
        writer = PdfWriter()
        writer.add_page(reader.pages[0])
        trimmed_path = output_dir / (docx_path.stem + "_trimmed.pdf")
        with open(trimmed_path, "wb") as f:
            writer.write(f)
        pdf_path.unlink()
        trimmed_path.rename(pdf_path)
        logger.info(f"Trimmed {len(reader.pages)}-page PDF to single page")

    return pdf_path

# Models
class Certificate(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    session_id: str
    participant_id: str
    certificate_number: str
    file_path: Optional[str] = None
    issue_date: datetime = Field(default_factory=get_malaysia_time)


router = APIRouter(prefix="/certificates", tags=["certificates"])


@router.get("/participant/{participant_id}")
async def get_participant_certificates(participant_id: str, current_user: User = Depends(get_current_user)):
    """Get certificates for a specific participant"""
    if current_user.role == "participant" and current_user.id != participant_id:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    certificates = await db.certificates.find({"participant_id": participant_id}, {"_id": 0}).to_list(100)
    for cert in certificates:
        if isinstance(cert.get('issue_date'), str):
            cert['issue_date'] = datetime.fromisoformat(cert['issue_date'])
    return certificates


@router.get("/my-certificates")
async def get_my_certificates(current_user: User = Depends(get_current_user)):
    """Get certificates for the current logged-in user (participant)"""
    if current_user.role != "participant":
        raise HTTPException(status_code=403, detail="Only participants can access this endpoint")
    
    certificates = await db.certificates.find({"participant_id": current_user.id}, {"_id": 0}).to_list(100)
    for cert in certificates:
        if isinstance(cert.get('issue_date'), str):
            cert['issue_date'] = datetime.fromisoformat(cert['issue_date'])
    return certificates


@router.get("/session/{session_id}")
async def get_session_certificates(session_id: str, current_user: User = Depends(get_current_user)):
    """Get all certificates for a session"""
    if current_user.role not in ["admin", "coordinator", "supervisor"]:
        raise HTTPException(status_code=403, detail="Only admins and coordinators can access certificates")
    
    access_records = await db.participant_access.find(
        {"session_id": session_id, "certificate_url": {"$exists": True, "$ne": None}},
        {"_id": 0}
    ).to_list(1000)
    
    certificates = []
    for access in access_records:
        if access.get('certificate_url'):
            certificates.append({
                "participant_id": access.get('participant_id'),
                "file_path": access.get('certificate_url'),
                "certificate_url": access.get('certificate_url'),
                "uploaded_at": access.get('certificate_uploaded_at'),
                "uploaded_by": access.get('certificate_uploaded_by')
            })
    
    return certificates


@router.get("/repository")
async def get_certificate_repository(current_user: User = Depends(get_current_user)):
    """Get all certificates (admin only)"""
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    certificates = await db.certificates.find({}, {"_id": 0}).to_list(1000)
    
    for cert in certificates:
        if isinstance(cert.get('issue_date'), str):
            cert['issue_date'] = datetime.fromisoformat(cert['issue_date'])
        
        if cert.get("participant_id"):
            user = await db.users.find_one({"id": cert["participant_id"]}, {"_id": 0, "full_name": 1})
            cert["participant_name"] = user.get("full_name") if user else "Unknown"
        
        if cert.get("session_id"):
            session = await db.sessions.find_one({"id": cert["session_id"]}, {"_id": 0, "name": 1})
            cert["session_name"] = session.get("name") if session else "Unknown"
    
    return certificates


@router.post("/upload/{session_id}/{participant_id}")
async def upload_participant_certificate(
    session_id: str,
    participant_id: str,
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user)
):
    """Upload certificate PDF for a specific participant"""
    if current_user.role == "coordinator":
        session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
        if not session:
            raise HTTPException(status_code=404, detail="Session not found")
        if session.get("coordinator_id") != current_user.id:
            raise HTTPException(status_code=403, detail="You can only upload certificates for your assigned sessions")
    elif current_user.role != "admin":
        raise HTTPException(status_code=403, detail="Only coordinators and admins can upload certificates")
    
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are accepted")
    
    settings = await db.settings.find_one({"id": "app_settings"}, {"_id": 0})
    max_size_mb = settings.get('max_certificate_file_size_mb', 5) if settings else 5
    max_size_bytes = max_size_mb * 1024 * 1024
    
    file.file.seek(0, 2)
    file_size = file.file.tell()
    file.file.seek(0)
    
    if file_size > max_size_bytes:
        raise HTTPException(status_code=400, detail=f"File size exceeds maximum allowed size of {max_size_mb}MB")
    
    unique_filename = f"{session_id}_{participant_id}_{uuid.uuid4().hex[:8]}.pdf"
    file_path = CERTIFICATE_PDF_DIR / unique_filename
    
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    certificate_url = f"/api/static/certificates_pdf/{unique_filename}"
    
    await db.participant_access.update_one(
        {"participant_id": participant_id, "session_id": session_id},
        {"$set": {
            "certificate_url": certificate_url,
            "certificate_uploaded_at": get_malaysia_time().isoformat(),
            "certificate_uploaded_by": current_user.id
        }},
        upsert=True
    )
    
    return {
        "certificate_url": certificate_url,
        "message": "Certificate uploaded successfully",
        "file_size_mb": round(file_size / (1024 * 1024), 2)
    }


@router.get("/eligibility/{session_id}/{participant_id}")
async def check_certificate_eligibility(session_id: str, participant_id: str, current_user: User = Depends(get_current_user)):
    """Check if participant is eligible for certificate"""
    if current_user.role not in ["admin", "super_admin", "coordinator", "assistant_admin"]:
        raise HTTPException(status_code=403, detail="Access denied")
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    access = await db.participant_access.find_one(
        {"participant_id": participant_id, "session_id": session_id},
        {"_id": 0}
    )
    
    attendance = await db.attendance.find_one(
        {"participant_id": participant_id, "session_id": session_id, "clock_out": {"$ne": None}},
        {"_id": 0}
    )
    
    post_test = await db.test_results.find_one(
        {"participant_id": participant_id, "session_id": session_id, "test_type": {"$in": ["post", "post_test"]}},
        {"_id": 0}
    )
    
    feedback = await db.course_feedback.find_one(
        {"participant_id": participant_id, "session_id": session_id},
        {"_id": 0}
    )
    
    return {
        "eligible": all([attendance, post_test and post_test.get("passed"), feedback]),
        "clocked_out": bool(attendance),
        "post_test_passed": bool(post_test and post_test.get("passed")),
        "feedback_submitted": bool(feedback),
        "certificate_uploaded": bool(access and access.get("certificate_url"))
    }


@router.post("/generate/{session_id}/{participant_id}")
async def generate_certificate(session_id: str, participant_id: str, current_user: User = Depends(get_current_user)):
    """Generate a certificate for a participant"""
    if current_user.role not in ["admin", "coordinator"]:
        raise HTTPException(status_code=403, detail="Only admins and coordinators can generate certificates")
    
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    participant = await db.users.find_one({"id": participant_id}, {"_id": 0})
    if not participant:
        raise HTTPException(status_code=404, detail="Participant not found")
    
    # Generate certificate number
    year = datetime.now().year
    count = await db.certificates.count_documents({"certificate_number": {"$regex": f"^CERT{year}"}})
    certificate_number = f"CERT{year}{str(count + 1).zfill(5)}"
    
    certificate = {
        "id": str(uuid.uuid4()),
        "session_id": session_id,
        "participant_id": participant_id,
        "certificate_number": certificate_number,
        "participant_name": participant.get("full_name"),
        "session_name": session.get("name"),
        "program_id": session.get("program_id"),
        "issue_date": get_malaysia_time().isoformat(),
        "generated_by": current_user.id
    }
    
    await db.certificates.insert_one(certificate)
    certificate.pop("_id", None)
    
    return {"message": "Certificate generated", "certificate": certificate}


@router.get("/download/{certificate_id}")
async def download_certificate_by_id(certificate_id: str, current_user: User = Depends(get_current_user)):
    """Download certificate by certificate ID"""
    certificate = await db.certificates.find_one({"id": certificate_id}, {"_id": 0})
    if not certificate:
        raise HTTPException(status_code=404, detail="Certificate not found")
    
    if current_user.role == "participant" and current_user.id != certificate.get("participant_id"):
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    file_path = certificate.get("file_path")
    if not file_path:
        raise HTTPException(status_code=404, detail="Certificate file not found")
    
    filename = file_path.split('/')[-1]
    full_path = CERTIFICATE_PDF_DIR / filename
    
    if not full_path.exists():
        raise HTTPException(status_code=404, detail="Certificate file not found on disk")
    
    return FileResponse(full_path, media_type="application/pdf", filename=f"certificate_{certificate_id}.pdf")


@router.get("/download/{session_id}/{participant_id}")
async def download_participant_certificate(session_id: str, participant_id: str, current_user: User = Depends(get_current_user)):
    """Download certificate for a participant in a session"""
    if current_user.id != participant_id and current_user.role not in ["admin", "coordinator"]:
        raise HTTPException(status_code=403, detail="Unauthorized")
    
    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    
    access = await db.participant_access.find_one(
        {"participant_id": participant_id, "session_id": session_id},
        {"_id": 0}
    )
    
    if not access or not access.get('certificate_url'):
        raise HTTPException(status_code=404, detail="No certificate uploaded for this participant")
    
    if current_user.id == participant_id:
        feedback_done = access.get('feedback_submitted', False) or access.get('feedback_completed', False)
        if not feedback_done:
            raise HTTPException(status_code=403, detail="Certificate not available. Please submit your feedback first.")
        
        attendance = await db.attendance.find_one(
            {"participant_id": participant_id, "session_id": session_id, "clock_out": {"$ne": None}},
            {"_id": 0}
        )
        if not attendance:
            raise HTTPException(status_code=403, detail="Certificate not available. Please clock out first.")
    
    filename = access['certificate_url'].split('/')[-1]
    file_path = CERTIFICATE_PDF_DIR / filename
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Certificate file not found")
    
    participant = await db.users.find_one({"id": participant_id}, {"_id": 0})
    participant_name = participant.get('full_name', 'participant').replace(' ', '_') if participant else 'participant'
    
    return FileResponse(file_path, media_type="application/pdf", filename=f"{participant_name}_certificate.pdf")


@router.get("/preview/{certificate_id}")
async def preview_certificate(certificate_id: str, current_user: User = Depends(get_current_user)):
    """Preview certificate (returns certificate data)"""
    if current_user.role not in ["admin", "super_admin", "coordinator", "assistant_admin", "supervisor"]:
        raise HTTPException(status_code=403, detail="Access denied")
    certificate = await db.certificates.find_one({"id": certificate_id}, {"_id": 0})
    if not certificate:
        raise HTTPException(status_code=404, detail="Certificate not found")
    
    if isinstance(certificate.get('issue_date'), str):
        certificate['issue_date'] = datetime.fromisoformat(certificate['issue_date'])
    
    return certificate


@router.get("/preview-image/{session_id}/{participant_id}")
async def preview_certificate_image(session_id: str, participant_id: str, current_user: User = Depends(get_current_user)):
    """Return the generated certificate as an inline PNG image for preview."""
    if current_user.role not in ["admin", "super_admin", "coordinator", "assistant_admin", "supervisor"]:
        raise HTTPException(status_code=403, detail="Access denied")

    access = await db.participant_access.find_one(
        {"participant_id": participant_id, "session_id": session_id}, {"_id": 0}
    )
    if not access or not access.get("certificate_url"):
        raise HTTPException(status_code=404, detail="No certificate found for this participant")

    filename = access["certificate_url"].split("/")[-1]
    pdf_path = CERTIFICATE_PDF_DIR / filename
    if not pdf_path.exists():
        raise HTTPException(status_code=404, detail="Certificate PDF file not found")

    # Convert first page to PNG
    from pdf2image import convert_from_path
    images = convert_from_path(str(pdf_path), dpi=150, first_page=1, last_page=1)
    import io
    buf = io.BytesIO()
    images[0].save(buf, format="PNG")
    buf.seek(0)

    from fastapi.responses import StreamingResponse
    return StreamingResponse(buf, media_type="image/png", headers={"Content-Disposition": "inline"})


# ==================== CERTIFICATE PDF GENERATION ====================

async def _check_eligibility(session_id: str, participant_id: str) -> dict:
    """Check if a participant meets certificate eligibility requirements."""
    attendance = await db.attendance.find_one(
        {"participant_id": participant_id, "session_id": session_id, "clock_out": {"$ne": None}},
        {"_id": 0}
    )
    post_test = await db.test_results.find_one(
        {"participant_id": participant_id, "session_id": session_id, "test_type": {"$in": ["post", "post_test"]}},
        {"_id": 0}
    )
    feedback = await db.course_feedback.find_one(
        {"participant_id": participant_id, "session_id": session_id},
        {"_id": 0}
    )
    return {
        "eligible": all([attendance, post_test and post_test.get("passed"), feedback]),
        "clocked_out": bool(attendance),
        "post_test_passed": bool(post_test and post_test.get("passed")),
        "feedback_submitted": bool(feedback),
    }


async def _generate_single_certificate_pdf(session: dict, participant: dict, current_user_id: str, font_settings: dict = None) -> dict:
    """Core logic: generate one certificate PDF from the .docx template."""
    template_path = TEMPLATE_DIR / "certificate_template.docx"
    if not template_path.exists():
        raise HTTPException(status_code=400, detail="Certificate template not uploaded. Go to Settings to upload.")

    if not font_settings:
        font_settings = await _get_font_settings()

    cert_number = await _generate_cert_number()
    replacements = await _build_replacements(session, participant, cert_number)

    doc = Document(str(template_path))
    _replace_placeholders_in_doc(doc, replacements, font_settings)

    # Apply margin and spacing from settings
    from docx.shared import Emu
    top_margin_pct = font_settings.get("top_margin", 80)
    spacing_pct = font_settings.get("paragraph_spacing", 65)
    original_top = 1231900  # original EMU value from template

    for section in doc.sections:
        section.top_margin = int(original_top * top_margin_pct / 100)

    for paragraph in doc.paragraphs:
        pf = paragraph.paragraph_format
        if pf.space_before and pf.space_before > Emu(50000):
            pf.space_before = int(pf.space_before * spacing_pct / 100)

    # Save temp docx, convert to PDF
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_docx = Path(tmp_dir) / f"cert_{participant['id']}.docx"
        doc.save(str(tmp_docx))

        pdf_path = _docx_to_pdf(tmp_docx, Path(tmp_dir))

        # Move PDF to permanent location
        final_filename = f"cert_{session['id']}_{participant['id']}_{uuid.uuid4().hex[:8]}.pdf"
        final_path = CERTIFICATE_PDF_DIR / final_filename
        shutil.move(str(pdf_path), str(final_path))

    certificate_url = f"/api/static/certificates_pdf/{final_filename}"

    # Save certificate record
    cert_record = {
        "id": str(uuid.uuid4()),
        "session_id": session["id"],
        "participant_id": participant["id"],
        "certificate_number": cert_number,
        "participant_name": participant.get("full_name"),
        "session_name": session.get("name"),
        "program_id": session.get("program_id"),
        "file_path": certificate_url,
        "issue_date": get_malaysia_time().isoformat(),
        "generated_by": current_user_id,
        "generation_method": "auto_template",
    }
    await db.certificates.insert_one(cert_record)
    cert_record.pop("_id", None)

    # Update participant_access with certificate URL
    await db.participant_access.update_one(
        {"participant_id": participant["id"], "session_id": session["id"]},
        {"$set": {
            "certificate_url": certificate_url,
            "certificate_uploaded_at": get_malaysia_time().isoformat(),
            "certificate_uploaded_by": current_user_id,
        }},
        upsert=True
    )

    return cert_record


@router.post("/generate-pdf/{session_id}/{participant_id}")
async def generate_certificate_pdf(
    session_id: str,
    participant_id: str,
    force: bool = Query(False, description="Skip eligibility checks"),
    current_user: User = Depends(get_current_user),
):
    """Generate a certificate PDF for a single participant from the .docx template."""
    if current_user.role not in ["admin", "coordinator"]:
        raise HTTPException(status_code=403, detail="Only admins and coordinators can generate certificates")

    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    participant = await db.users.find_one({"id": participant_id}, {"_id": 0})
    if not participant:
        raise HTTPException(status_code=404, detail="Participant not found")

    if not force:
        eligibility = await _check_eligibility(session_id, participant_id)
        if not eligibility["eligible"]:
            reasons = []
            if not eligibility["clocked_out"]:
                reasons.append("attendance not completed")
            if not eligibility["post_test_passed"]:
                reasons.append("post-test not passed")
            if not eligibility["feedback_submitted"]:
                reasons.append("feedback not submitted")
            raise HTTPException(
                status_code=400,
                detail=f"Participant not eligible: {', '.join(reasons)}. Use force=true to override."
            )

    cert = await _generate_single_certificate_pdf(session, participant, current_user.id)
    return {"message": "Certificate generated successfully", "certificate": cert}


@router.post("/generate-bulk-pdf/{session_id}")
async def generate_bulk_certificate_pdf(
    session_id: str,
    force: bool = Query(False, description="Skip eligibility checks"),
    current_user: User = Depends(get_current_user),
):
    """Generate certificate PDFs for all (eligible) participants in a session."""
    if current_user.role not in ["admin", "coordinator"]:
        raise HTTPException(status_code=403, detail="Only admins and coordinators can generate certificates")

    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    participant_ids = session.get("participant_ids", [])
    if not participant_ids:
        raise HTTPException(status_code=400, detail="No participants in this session")

    results = {"generated": [], "skipped": [], "errors": []}

    for pid in participant_ids:
        participant = await db.users.find_one({"id": pid}, {"_id": 0})
        if not participant:
            results["errors"].append({"participant_id": pid, "reason": "Participant not found"})
            continue

        if not force:
            eligibility = await _check_eligibility(session_id, pid)
            if not eligibility["eligible"]:
                results["skipped"].append({
                    "participant_id": pid,
                    "name": participant.get("full_name", "Unknown"),
                    "reason": eligibility,
                })
                continue

        try:
            cert = await _generate_single_certificate_pdf(session, participant, current_user.id)
            results["generated"].append({
                "participant_id": pid,
                "name": participant.get("full_name"),
                "certificate_number": cert["certificate_number"],
            })
        except Exception as e:
            logger.error(f"Failed to generate cert for {pid}: {e}")
            results["errors"].append({
                "participant_id": pid,
                "name": participant.get("full_name", "Unknown"),
                "reason": str(e),
            })

    return {
        "message": f"Generated {len(results['generated'])} certificates, skipped {len(results['skipped'])}, errors {len(results['errors'])}",
        "results": results,
    }


# ==================== FONT SETTINGS & PREVIEW ====================

@router.get("/font-settings")
async def get_font_settings(current_user: User = Depends(get_current_user)):
    """Get current certificate font settings."""
    if current_user.role not in ["admin", "coordinator"]:
        raise HTTPException(status_code=403, detail="Access denied")
    settings = await _get_font_settings()
    return settings


@router.put("/font-settings")
async def save_font_settings(settings: dict, current_user: User = Depends(get_current_user)):
    """Save certificate font settings."""
    if current_user.role not in ["admin", "coordinator"]:
        raise HTTPException(status_code=403, detail="Access denied")
    settings["id"] = "certificate_font_settings"
    await db.settings.update_one(
        {"id": "certificate_font_settings"},
        {"$set": settings},
        upsert=True,
    )
    return {"message": "Font settings saved"}


@router.post("/preview-pdf/{session_id}/{participant_id}")
async def preview_certificate_pdf(
    session_id: str,
    participant_id: str,
    font_settings: dict = None,
    current_user: User = Depends(get_current_user),
):
    """Generate a certificate preview as PNG using provided font settings (no DB record)."""
    if current_user.role not in ["admin", "coordinator"]:
        raise HTTPException(status_code=403, detail="Access denied")

    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    participant = await db.users.find_one({"id": participant_id}, {"_id": 0})
    if not participant:
        raise HTTPException(status_code=404, detail="Participant not found")

    template_path = TEMPLATE_DIR / "certificate_template.docx"
    if not template_path.exists():
        raise HTTPException(status_code=400, detail="Certificate template not uploaded")

    # Use provided settings or load saved ones
    if not font_settings:
        font_settings = await _get_font_settings()

    # Build replacements with a preview cert number
    replacements = await _build_replacements(session, participant, "MDDRC/COA/PREVIEW/00000")

    doc = Document(str(template_path))
    _replace_placeholders_in_doc(doc, replacements, font_settings)

    # Apply margin and spacing
    from docx.shared import Emu
    top_margin_pct = font_settings.get("top_margin", 80)
    spacing_pct = font_settings.get("paragraph_spacing", 65)
    original_top = 1231900

    for section in doc.sections:
        section.top_margin = int(original_top * top_margin_pct / 100)

    for paragraph in doc.paragraphs:
        pf = paragraph.paragraph_format
        if pf.space_before and pf.space_before > Emu(50000):
            pf.space_before = int(pf.space_before * spacing_pct / 100)

    # Save temp docx, convert to PDF, then to PNG
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_docx = Path(tmp_dir) / "preview.docx"
        doc.save(str(tmp_docx))
        pdf_path = _docx_to_pdf(tmp_docx, Path(tmp_dir))

        from pdf2image import convert_from_path
        images = convert_from_path(str(pdf_path), dpi=150, first_page=1, last_page=1)

        import io
        buf = io.BytesIO()
        images[0].save(buf, format="PNG")
        buf.seek(0)

    from fastapi.responses import StreamingResponse
    return StreamingResponse(buf, media_type="image/png", headers={"Content-Disposition": "inline"})


# ==================== VISUAL DESIGNER ENDPOINTS ====================

@router.get("/designer-layout")
async def get_designer_layout(current_user: User = Depends(get_current_user)):
    """Get saved certificate designer layout."""
    if current_user.role not in ["admin", "coordinator"]:
        raise HTTPException(status_code=403, detail="Access denied")
    saved = await db.settings.find_one({"id": "certificate_designer_layout"}, {"_id": 0})
    if saved:
        saved.pop("id", None)
        return saved
    return {}


@router.put("/designer-layout")
async def save_designer_layout(layout: dict, current_user: User = Depends(get_current_user)):
    """Save certificate designer layout."""
    if current_user.role not in ["admin", "coordinator"]:
        raise HTTPException(status_code=403, detail="Access denied")
    layout["id"] = "certificate_designer_layout"
    await db.settings.update_one(
        {"id": "certificate_designer_layout"},
        {"$set": layout},
        upsert=True,
    )
    return {"message": "Layout saved"}


@router.post("/preview-data")
async def get_preview_data(body: dict, current_user: User = Depends(get_current_user)):
    """Get live participant data formatted for the designer preview."""
    if current_user.role not in ["admin", "coordinator"]:
        raise HTTPException(status_code=403, detail="Access denied")

    session_id = body.get("session_id")
    participant_id = body.get("participant_id")

    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    participant = await db.users.find_one({"id": participant_id}, {"_id": 0})
    if not participant:
        raise HTTPException(status_code=404, detail="Participant not found")

    replacements = await _build_replacements(session, participant, "MDDRC/COA/PREVIEW/00000")
    same_date = replacements.pop("_same_date", False)

    start = replacements.get("{{TRAINING_DATE}}", "")
    end = replacements.get("{{END_DATE}}", "")
    dates = start if same_date or not end else f"{start} - {end}"

    validity_start = replacements.get("{{VALIDITY_START}}", "")
    validity_end = replacements.get("{{VALIDITY_END}}", "")
    validity = f"Valid: {validity_start} - {validity_end}" if validity_end else ""

    return {
        "participant_name": replacements.get("{{PARTICIPANT_NAME}}", ""),
        "ic_number": f"I.C. No: {replacements.get('{{IC_NUMBER}}', '')}",
        "company_name": replacements.get("{{COMPANY_NAME}}", ""),
        "certificate_title": replacements.get("{{CERTIFICATE_TITLE}}", ""),
        "certificate_subtitle": replacements.get("{{CERTIFICATE_SUBTITLE}}", ""),
        "training_dates": dates,
        "venue": replacements.get("{{VENUE}}", ""),
        "validity": validity,
        "certificate_number": f"Certificate Serial No: {replacements.get('{{CERTIFICATE_NUMBER}}', '')}",
    }


async def _get_designer_layout() -> dict:
    """Load designer layout from DB."""
    saved = await db.settings.find_one({"id": "certificate_designer_layout"}, {"_id": 0})
    if saved:
        saved.pop("id", None)
        return saved
    return {}


def _build_certificate_html(layout: dict, data: dict, bg_path: str) -> str:
    """Build HTML string for certificate PDF generation."""
    # A4 dimensions at 96 DPI: 794 x 1123 px
    W, H = 794, 1123

    elements = ""
    for key, field in layout.items():
        if not isinstance(field, dict) or "x" not in field:
            continue
        text = data.get(key, "")
        if not text:
            continue
        left = field.get("x", 0) * W / 100
        top = field.get("y", 0) * H / 100
        width = field.get("width", 50) * W / 100
        font_size = field.get("fontSize", 12)
        font_weight = field.get("fontWeight", "normal")
        text_align = field.get("textAlign", "center")
        color = field.get("color", "#000000")

        elements += f"""<div style="
            position:absolute; left:{left:.1f}px; top:{top:.1f}px; width:{width:.1f}px;
            font-size:{font_size}pt; font-weight:{font_weight}; text-align:{text_align};
            color:{color}; font-family:Calibri,Arial,sans-serif; line-height:1.2;
            white-space:nowrap; overflow:visible;
        ">{text}</div>\n"""

    return f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
    @page {{ size: A4 portrait; margin: 0; }}
    * {{ margin:0; padding:0; box-sizing:border-box; }}
    body {{ width:794px; height:1123px; position:relative; overflow:hidden; }}
</style></head>
<body>
    <img src="{bg_path}" style="width:794px;height:1123px;position:absolute;top:0;left:0;" />
    {elements}
</body></html>"""


@router.post("/generate-designed/{session_id}/{participant_id}")
async def generate_designed_certificate(
    session_id: str,
    participant_id: str,
    force: bool = Query(False),
    current_user: User = Depends(get_current_user),
):
    """Generate certificate PDF using the visual designer layout (single participant)."""
    if current_user.role not in ["admin", "coordinator"]:
        raise HTTPException(status_code=403, detail="Access denied")

    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    participant = await db.users.find_one({"id": participant_id}, {"_id": 0})
    if not participant:
        raise HTTPException(status_code=404, detail="Participant not found")

    if not force:
        eligibility = await _check_eligibility(session_id, participant_id)
        if not eligibility["eligible"]:
            reasons = []
            if not eligibility["clocked_out"]:
                reasons.append("attendance not completed")
            if not eligibility["post_test_passed"]:
                reasons.append("post-test not passed")
            if not eligibility["feedback_submitted"]:
                reasons.append("feedback not submitted")
            raise HTTPException(status_code=400, detail=f"Not eligible: {', '.join(reasons)}. Use force=true to override.")

    layout = await _get_designer_layout()
    cert_number = await _generate_cert_number()
    replacements = await _build_replacements(session, participant, cert_number)
    same_date = replacements.pop("_same_date", False)

    start = replacements.get("{{TRAINING_DATE}}", "")
    end = replacements.get("{{END_DATE}}", "")
    data = {
        "participant_name": replacements.get("{{PARTICIPANT_NAME}}", ""),
        "ic_number": f"I.C. No: {replacements.get('{{IC_NUMBER}}', '')}",
        "company_name": replacements.get("{{COMPANY_NAME}}", ""),
        "certificate_title": replacements.get("{{CERTIFICATE_TITLE}}", ""),
        "certificate_subtitle": replacements.get("{{CERTIFICATE_SUBTITLE}}", ""),
        "training_dates": start if same_date or not end else f"{start} - {end}",
        "venue": replacements.get("{{VENUE}}", ""),
        "validity": f"Valid: {replacements.get('{{VALIDITY_START}}', '')} - {replacements.get('{{VALIDITY_END}}', '')}" if replacements.get("{{VALIDITY_END}}") else "",
        "certificate_number": f"Certificate Serial No: {cert_number}",
    }

    # Build HTML and save
    bg_abs_path = str(TEMPLATE_DIR / "cert_background.png")
    html_content = _build_certificate_html(layout, data, f"file://{bg_abs_path}")

    # Save HTML for client-side PDF generation (return HTML for frontend to convert)
    # OR generate server-side with wkhtmltopdf if available
    # For now, return the data for frontend PDF generation
    certificate_url = ""

    cert_record = {
        "id": str(uuid.uuid4()),
        "session_id": session["id"],
        "participant_id": participant["id"],
        "certificate_number": cert_number,
        "participant_name": participant.get("full_name"),
        "session_name": session.get("name"),
        "program_id": session.get("program_id"),
        "file_path": certificate_url,
        "issue_date": get_malaysia_time().isoformat(),
        "generated_by": current_user.id,
        "generation_method": "designer",
    }
    await db.certificates.insert_one(cert_record)
    cert_record.pop("_id", None)

    await db.participant_access.update_one(
        {"participant_id": participant["id"], "session_id": session["id"]},
        {"$set": {
            "certificate_uploaded_at": get_malaysia_time().isoformat(),
            "certificate_uploaded_by": current_user.id,
        }},
        upsert=True
    )

    return {
        "message": "Certificate generated successfully",
        "certificate": cert_record,
        "layout": layout,
        "data": data,
        "bg_url": "/api/static/templates/cert_background.png",
    }


@router.post("/generate-designed/{session_id}")
async def generate_designed_bulk(
    session_id: str,
    force: bool = Query(False),
    current_user: User = Depends(get_current_user),
):
    """Bulk generate certificates using the designer layout."""
    if current_user.role not in ["admin", "coordinator"]:
        raise HTTPException(status_code=403, detail="Access denied")

    session = await db.sessions.find_one({"id": session_id}, {"_id": 0})
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    pids = session.get("participant_ids", [])
    if not pids:
        raise HTTPException(status_code=400, detail="No participants")

    layout = await _get_designer_layout()
    results = {"generated": [], "skipped": [], "errors": []}

    for pid in pids:
        participant = await db.users.find_one({"id": pid}, {"_id": 0})
        if not participant:
            results["errors"].append({"participant_id": pid, "reason": "Not found"})
            continue

        if not force:
            elig = await _check_eligibility(session_id, pid)
            if not elig["eligible"]:
                results["skipped"].append({"participant_id": pid, "name": participant.get("full_name"), "reason": elig})
                continue

        try:
            cert_number = await _generate_cert_number()
            replacements = await _build_replacements(session, participant, cert_number)
            same_date = replacements.pop("_same_date", False)
            start = replacements.get("{{TRAINING_DATE}}", "")
            end = replacements.get("{{END_DATE}}", "")

            cert_record = {
                "id": str(uuid.uuid4()),
                "session_id": session["id"],
                "participant_id": pid,
                "certificate_number": cert_number,
                "participant_name": participant.get("full_name"),
                "session_name": session.get("name"),
                "program_id": session.get("program_id"),
                "file_path": "",
                "issue_date": get_malaysia_time().isoformat(),
                "generated_by": current_user.id,
                "generation_method": "designer",
            }
            await db.certificates.insert_one(cert_record)
            cert_record.pop("_id", None)

            results["generated"].append({
                "participant_id": pid,
                "name": participant.get("full_name"),
                "certificate_number": cert_number,
            })
        except Exception as e:
            logger.error(f"Cert gen failed for {pid}: {e}")
            results["errors"].append({"participant_id": pid, "name": participant.get("full_name", "?"), "reason": str(e)})

    return {
        "message": f"Generated {len(results['generated'])}, skipped {len(results['skipped'])}, errors {len(results['errors'])}",
        "results": results,
    }
